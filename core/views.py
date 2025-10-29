from django.shortcuts import render
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from docx import Document
import io, re

def index(request):
    return render(request, 'index.html')

def upload_docx(request):
    if request.method == 'POST' and request.FILES['document']:
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        file_path = fs.path(filename)

        # Lê o documento e detecta placeholders
        doc = Document(file_path)
        placeholders = set()

        for para in doc.paragraphs:
            matches = re.findall(r'\{([^{}]+)\}', para.text)
            placeholders.update(matches)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{([^{}]+)\}', cell.text)
                    placeholders.update(matches)

        return render(request, 'fill_fields.html', {
            'filename': filename,
            'placeholders': sorted(list(placeholders))
        })

    return HttpResponse("Nenhum arquivo enviado.")

def generate_docx(request):
    if request.method == 'POST':
        filename = request.POST.get('filename')
        if not filename:
            return HttpResponse("Arquivo não encontrado no formulário.")

        fs = FileSystemStorage()
        file_path = fs.path(filename)
        doc = Document(file_path)

        # Função que substitui placeholders mesmo que estejam quebrados em runs
        def replace_placeholders_in_paragraph(paragraph):
            # Junta todo o texto do parágrafo
            full_text = ''.join(run.text for run in paragraph.runs)
            
            for key, value in request.POST.items():
                if key == 'filename':
                    continue
                # Substitui com regex, ignorando espaços dentro das chaves
                pattern = re.compile(r'\{\s*' + re.escape(key) + r'\s*\}')
                full_text = pattern.sub(value, full_text)
            
            # Atualiza o primeiro run e limpa os outros
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ''

        # Substitui em parágrafos fora de tabelas
        for para in doc.paragraphs:
            replace_placeholders_in_paragraph(para)

        # Substitui em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_paragraph(para)

        # Salva no buffer para download
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=documento_preenchido.docx'
        return response

    return HttpResponse("Requisição inválida.")
